from flask import Blueprint,render_template, request,redirect, url_for, flash, get_flashed_messages,current_app,send_from_directory
from flask_login import login_required
from Rocho_app import administrador
from Rocho_app.pacientes.model.paciente import Paciente
from Rocho_app.pacientes.model.categoria import Categoria
from Rocho_app.pacientes.model.paciente import PacienteForm
from docx import Document
from docx.shared import Inches
from Rocho_app import descargas
import os
pacientes = Blueprint('pacientes',__name__)
@pacientes.before_request
@login_required
def administrador():
   pass

@pacientes.route('/home')
@pacientes.route('/home/<int:page>')
def home(page=1):
   print()
   return render_template('index.html',pacientes=Paciente.query.paginate(page,5))

@pacientes.route('/paciente/<int:id>')
def paciente(id):
   paciente = Paciente.query.get_or_404(id)
   return render_template('paciente.html',paciente=paciente)

@pacientes.route('/paciente-delete/<int:id>')
def delete(id):
   paciente = Paciente.query.get_or_404(id)
   db.session.delete(paciente) #borrar un registro a la base de datos
   db.session.commit()
   flash("Paciente eliminado con exito")
   return redirect(url_for('pacientes.home'))

@pacientes.route('/paciente-crear',methods=['GET','POST'])
def crear():
   form = PacienteForm(meta={'csrf':False})
   categorias =[(c.id,c.nombre) for c in Categoria.query.all() ]
   form.categoria_id.choices = categorias
   if form.validate_on_submit():
      p = Paciente(request.form['nombre'],request.form['apellidos'],request.form['edad'],request.form['peso'],request.form['talla'],request.form['grupo_s'],request.form['temperatura'],request.form['fecha'],request.form['sintomas'],request.form['padecimientos'],request.form['alergias'],request.form['tratamiento'],request.form['categoria_id']) 
      db.session.add(p) #agregar un registro a la base de datos
      db.session.commit() #agregar instruccion para que se pueda subir a la base de datos 
      flash("Paciente creado con exito")
      return redirect(url_for('pacientes.crear'))
   if form.errors:
      flash(form.errors,'danger')

   return render_template('crear.html',form=form)


@pacientes.route('/paciente-update/<int:id>',methods=['GET','POST'])
def update(id):
   paciente = Paciente.query.get_or_404(id)
   form = PacienteForm(meta={'csrf':False})
   categorias =[(c.id,c.nombre) for c in Categoria.query.all() ]
   form.categoria_id.choices = categorias
   if request.method == 'GET':
      form.nombre.data=paciente.nombre
      form.apellidos.data=paciente.apellidos
      form.edad.data=paciente.edad
      form.peso.data=paciente.peso
      form.talla.data=paciente.talla
      form.grupo_s.data=paciente.grupo_s
      form.temperatura.data=paciente.temperatura
      form.fecha.data=paciente.fecha
      form.sintomas.data=paciente.sintomas
      form.padecimientos.data=paciente.padecimientos
      form.alergias.data=paciente.alergias
      form.tratamiento.data=paciente.tratamiento 
      form.categoria_id.data = paciente.categoria_id
   if form.validate_on_submit():

      paciente.nombre = form.nombre.data
      paciente.apellidos = form.apellidos.data
      paciente.edad = form.edad.data
      paciente.peso=form.peso.data
      paciente.talla=form.talla.data
      paciente.grupo_s=form.grupo_s.data
      paciente.temperatura=form.temperatura.data
      paciente.fecha=form.fecha.data
      paciente.sintomas = form.sintomas.data
      paciente.padecimientos = form.padecimientos.data
      paciente.alergias = form.alergias.data
      paciente.tratamiento = form.tratamiento.data
      paciente.categoria_id = form.categoria_id.data
      db.session.add(paciente) #agregar un registro a la base de datos
      db.session.commit() #agregar instruccion para que se pueda subir a la base de datos 
      flash("Paciente actualizado con exito")
      return redirect(url_for('pacientes.update',id=paciente.id))
      if form.errors:
         flash(form.errors,'danger')
   return render_template('update.html',paciente=paciente, form=form)

@pacientes.route('/receta/<int:id>',methods=['GET','POST'])
def receta(id):
   paciente = Paciente.query.get_or_404(id)
   form = PacienteForm(meta={'csrf':False})
   categorias =[(c.id,c.nombre) for c in Categoria.query.all() ]
   form.categoria_id.choices = categorias
   if request.method == 'GET':
      form.nombre.data=paciente.nombre
      form.apellidos.data=paciente.apellidos
      form.edad.data=paciente.edad
      form.peso.data=paciente.peso
      form.talla.data=paciente.talla
      form.grupo_s.data=paciente.grupo_s
      form.temperatura.data=paciente.temperatura
      form.fecha.data=paciente.fecha
      form.sintomas.data=paciente.sintomas
      form.padecimientos.data=paciente.padecimientos
      form.alergias.data=paciente.alergias
      form.tratamiento.data=paciente.tratamiento 
      form.categoria_id.data = paciente.categoria_id
   if form.validate_on_submit():

      paciente.nombre = form.nombre.data
      paciente.apellidos = form.apellidos.data
      paciente.edad = form.edad.data
      paciente.peso=form.peso.data
      paciente.talla=form.talla.data
      paciente.grupo_s=form.grupo_s.data
      paciente.temperatura=form.temperatura.data
      paciente.fecha=form.fecha.data
      paciente.sintomas = form.sintomas.data
      paciente.padecimientos = form.padecimientos.data
      paciente.alergias = form.alergias.data
      paciente.tratamiento = form.tratamiento.data
      paciente.categoria_id = form.categoria_id.data
      #crear receta
      document = Document()
      document.add_heading("Hospital Mexico Inn", level=0) 
    
      document.add_heading('Dr. Edaurdo Perez Castañeda, Medicina General', level=1)
      document.add_paragraph('Ced.Prof:6456761757, Universidad Automa Metropolitana', style='Intense Quote')
      #document.add_heading(paciente.nombre+" "+paciente.apellidos, 0) 
      string_edad=str(paciente.edad)
      string_peso=str(paciente.peso)
      p = document.add_paragraph()
      p.add_run("Nombre: ").bold = True
      p.add_run(paciente.nombre+" "+paciente.apellidos)
      p.add_run("    ")
      p.add_run("Fecha: ").bold = True
      p.add_run(paciente.fecha)
      p.add_run("    ")
      p=document.add_paragraph()
      p.add_run("Edad: ").bold = True
      p.add_run(string_edad[0:2]+" años")
      p.add_run("    ")
      p.add_run("Peso: ").bold = True  
      p.add_run(string_peso[0:2]+" Kg")
      p.add_run("    ")
      p.add_run("Talla: ").bold = True  
      p.add_run(paciente.talla)
      p.add_run("    ")
      p.add_run("Grupo Sanguineo: ").bold = True  
      p.add_run(paciente.grupo_s)
      p.add_run("    ")
      p.add_run("Temperatura: ").bold = True  
      p.add_run(paciente.temperatura)
      p.add_run("    ")
      document.add_heading('Alergias:', level=1)
      document.add_paragraph(paciente.padecimientos)
      document.add_heading('Sintomas:', level=1)
      document.add_paragraph(paciente.sintomas)
      document.add_heading('Tratamiento:', level=1)
      document.add_paragraph(paciente.tratamiento)
      document.save('recetas/'+paciente.nombre+'.docx')
      flash("Receta generada con éxito")
      return redirect(url_for('pacientes.receta',id=paciente.id))
      if form.errors:
         flash(form.errors,'danger')
   return render_template('paciente.html',paciente=paciente, form=form)
@pacientes.route('/descargar/<path:filename>', methods=['GET', 'POST'])
def descargar(filename):
   descargar = os.path.join(current_app.root_path, descargas)
   return send_from_directory(directory=descargar, filename=filename)